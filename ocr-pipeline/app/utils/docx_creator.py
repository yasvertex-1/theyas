import os
import json
from typing import List, Dict, Any
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_log_file_with_metadata(extracted_pages: List[Dict[str, Any]], input_pdf_path: str, output_path: str, timing_info: Dict[str, float] = None) -> str:
    """
    Create a log file with document metadata and validation information.

    Args:
        extracted_pages: List of dictionaries with page data
        input_pdf_path: Original PDF path
        output_path: Path where to save the DOCX file (used to determine log location)
        timing_info: Optional dictionary with timing information

    Returns:
        Path to the created log file
    """
    # Create log filename based on DOCX filename
    log_path = output_path.replace('.docx', '_metadata.log')
    
    # Prepare metadata
    successful_extractions = sum(1 for page in extracted_pages if page.get('error') is None and page.get('text', '').strip())
    failed_extractions = len(extracted_pages) - successful_extractions
    
    metadata_info = {
        'Source PDF': os.path.basename(input_pdf_path),
        'Source Path': input_pdf_path,
        'Extraction Date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'Total Pages': len(extracted_pages),
        'Successful Extractions': successful_extractions,
        'Failed Extractions': failed_extractions
    }
    
    # Add timing information if available
    if timing_info:
        metadata_info.update({
            'Total Processing Time (seconds)': round(timing_info.get('total_time_seconds', 0), 2),
            'Text Extraction Time (seconds)': round(timing_info.get('extraction_time_seconds', 0), 2),
            'DOCX Creation Time (seconds)': round(timing_info.get('docx_creation_time_seconds', 0), 2),
            'Average Time per Page (seconds)': round(timing_info.get('average_time_per_page_seconds', 0), 2)
        })
    
    # Collect validation issues and page errors
    validation_summary = {
        'pages_with_issues': []
    }
    
    for page in extracted_pages:
        page_num = page.get('page_number', 0)
        validation_issues = page.get('validation_issues', [])
        error = page.get('error')
        
        if validation_issues or error:
            page_info = {'page_number': page_num}
            if error:
                page_info['error'] = str(error)
            if validation_issues:
                page_info['validation_issues'] = validation_issues
            validation_summary['pages_with_issues'].append(page_info)
    
    # Write log file
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("DOCUMENT EXTRACTION METADATA\n")
        f.write("=" * 80 + "\n\n")
        
        for key, value in metadata_info.items():
            f.write(f"{key}: {value}\n")
        
        if validation_summary['pages_with_issues']:
            f.write("\n" + "=" * 80 + "\n")
            f.write("VALIDATION ISSUES\n")
            f.write("=" * 80 + "\n\n")
            f.write(json.dumps(validation_summary, indent=2, ensure_ascii=False))
        
        f.write("\n")
    
    return log_path

def create_docx_from_extracted_pages(extracted_pages: List[Dict[str, Any]], input_pdf_path: str, output_path: str, timing_info: Dict[str, float] = None) -> None:
    """
    Create a DOCX file from extracted text pages.

    Args:
        extracted_pages: List of dictionaries with 'page_number', 'text', and 'error' keys
        input_pdf_path: Original PDF path for metadata
        output_path: Path where to save the DOCX file
        timing_info: Optional dictionary with timing information
    """
    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # Create metadata log file
    create_log_file_with_metadata(extracted_pages, input_pdf_path, output_path, timing_info)

    # Sort pages: first by page number, then by column (right before left)
    def sort_key(page):
        page_num = page.get('page_number', 0)
        column_type = page.get('column_type')
        # Right column = 0, Left column = 1, None = 0.5 (single page, in the middle)
        column_order = {'right': 0, 'left': 1, None: 0.5}
        return (page_num, column_order.get(column_type, 0.5))
    
    sorted_pages = sorted(extracted_pages, key=sort_key)

    # Create a new Document
    doc = Document()
    
    for page in sorted_pages:
        page_num = page.get('page_number', 0)
        column_type = page.get('column_type')
        text = page.get('text', '')
        structured_table = page.get('structured_table') or None
        notes = page.get('notes')
        
        # Add page header with column information
        heading_parts = [f'Page {page_num}']
        if column_type in ('left', 'right'):
            heading_parts.append(column_type.capitalize())
        doc.add_heading(' '.join(heading_parts), level=2)
        
        if structured_table:
            headers = structured_table.get('headers') or []
            rows = structured_table.get('rows') or []
            notes_text = structured_table.get('notes') or ''

            if headers:
                table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                table.style = 'Table Grid'
                for col_idx, header in enumerate(headers):
                    table.cell(0, col_idx).text = header
                start_row = 1
            else:
                max_cols = max((len(row) for row in rows), default=0)
                table = doc.add_table(rows=len(rows), cols=max(1, max_cols))
                table.style = 'Table Grid'
                start_row = 0

            for row_idx, row in enumerate(rows):
                for col_idx, cell_text in enumerate(row):
                    table.cell(row_idx + start_row, col_idx).text = cell_text or ''

            if notes_text:
                notes_para = doc.add_paragraph(notes_text)
                notes_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif text and text.strip():
            text_para = doc.add_paragraph(text.strip())
            text_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add page break except for the last page
        if page != sorted_pages[-1]:
            doc.add_page_break()
    
    # Save the document
    doc.save(output_path)

def create_output_docx_filename(input_filename: str) -> str:
    """
    Create output DOCX filename from input PDF filename.

    Args:
        input_filename: Original PDF filename

    Returns:
        Output DOCX filename
    """
    base_name = os.path.splitext(input_filename)[0]
    return f"{base_name}_extracted.docx"
